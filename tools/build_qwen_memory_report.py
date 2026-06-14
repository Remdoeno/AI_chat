import json
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
METRICS_LOG = ROOT / "log" / "qwen_memory_doc_metrics_20260528.log"
OUT = ROOT / "output" / "Qwen3带记忆大模型系统说明.docx"


def load_metrics() -> dict:
    text = METRICS_LOG.read_text(encoding="utf-8")
    start = text.find("{")
    end = text.rfind("}")
    if start < 0 or end < 0:
        raise RuntimeError(f"cannot parse metrics json from {METRICS_LOG}")
    return json.loads(text[start : end + 1])


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(table, top=80, start=120, bottom=80, end=120) -> None:
    tbl_pr = table._tbl.tblPr
    margins = tbl_pr.first_child_found_in("w:tblCellMar")
    if margins is None:
        margins = OxmlElement("w:tblCellMar")
        tbl_pr.append(margins)
    for name, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = margins.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            margins.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    for row in table.rows:
        for idx, width in enumerate(widths):
            cell = row.cells[idx]
            cell.width = Inches(width)
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(int(width * 1440)))
            tc_w.set(qn("w:type"), "dxa")


def set_run_font(run, east_asia="Microsoft YaHei") -> None:
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), east_asia)


def add_paragraph(doc, text="", style=None, bold_prefix=None):
    p = doc.add_paragraph(style=style)
    if bold_prefix and text.startswith(bold_prefix):
        first = p.add_run(bold_prefix)
        first.bold = True
        set_run_font(first)
        rest = p.add_run(text[len(bold_prefix) :])
        set_run_font(rest)
    else:
        run = p.add_run(text)
        set_run_font(run)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        run = p.add_run(item)
        set_run_font(run)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        run = p.add_run(item)
        set_run_font(run)


def style_table(table, header_fill="E8EEF5"):
    table.style = "Table Grid"
    set_cell_margins(table)
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.space_after = Pt(2)
                for run in paragraph.runs:
                    set_run_font(run)
                    run.font.size = Pt(9.5)
        if row_idx == 0:
            for cell in row.cells:
                set_cell_shading(cell, header_fill)
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.bold = True


def add_kv_table(doc, rows, widths=(2.0, 4.5)):
    table = doc.add_table(rows=1, cols=2)
    hdr = table.rows[0].cells
    hdr[0].text = "项目"
    hdr[1].text = "当前值 / 说明"
    for key, value in rows:
        cells = table.add_row().cells
        cells[0].text = str(key)
        cells[1].text = str(value)
    style_table(table)
    set_table_widths(table, widths)
    return table


def add_matrix_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = str(value)
    style_table(table)
    set_table_widths(table, widths)
    return table


def setup_styles(doc):
    section = doc.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for name in ("List Bullet", "List Number"):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.167


def add_header_footer(doc):
    section = doc.sections[0]
    header = section.header.paragraphs[0]
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = header.add_run("Qwen3 记忆系统说明")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(85, 85, 85)
    set_run_font(run)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = footer.add_run("生成时间：" + datetime.now().strftime("%Y-%m-%d %H:%M"))
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(85, 85, 85)
    set_run_font(run)


def build_doc():
    metrics = load_metrics()
    doc = Document()
    setup_styles(doc)
    add_header_footer(doc)

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(3)
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = title.add_run("Qwen3 带记忆大模型系统说明")
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = RGBColor(11, 37, 69)
    set_run_font(run)

    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(12)
    run = subtitle.add_run("基于 FastAPI、SQLite、Qwen3-Embedding-8B 与 Qwen3.6-35B-A3B 的网页聊天记忆系统")
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(85, 85, 85)
    set_run_font(run)

    doc.add_heading("1. 当前结论", level=1)
    add_paragraph(
        doc,
        "当前系统已经从关键词触发式记忆升级为 embedding 相似度检索式记忆。每轮对话仍按浏览器会话即时聊天，但历史消息会被软保留在 SQLite 中，并切分为连续聊天片段生成向量；后续用户提问时，系统用当前问题的向量去匹配历史片段，将最相关的片段注入 system prompt，形成跨会话记忆。",
    )
    add_bullets(
        doc,
        [
            "主模型负责生成回答：qwen3.6-35b-a3b-262k，OpenAI-compatible API 端口为 8000。",
            "记忆模型负责生成向量：qwen3-embedding-8b，OpenAI-compatible embeddings API 端口为 8001。",
            "记忆向量库保存在同一个 SQLite 数据库中，当前使用 NumPy 暴力余弦检索，规模较小时简单可靠。",
            "如果 embedding 服务不可用，后端会记录错误事件，并回退到原先的关键词/FTS 记忆逻辑。",
        ],
    )

    doc.add_heading("2. 系统组成", level=1)
    add_matrix_table(
        doc,
        ["组件", "位置 / 接口", "职责"],
        [
            ["网页后端", "/base/home/lizhzh/Project3/qwen_web/app.py", "FastAPI 服务、会话管理、流式聊天、记忆注入、增量向量刷新。"],
            ["聊天数据库", "qwen_web/data/chat_history.sqlite3", "保存 sessions、messages、events、memory_segments、memory_vectors 等表。"],
            ["主模型服务", "http://127.0.0.1:8000/v1", "调用 qwen3.6-35b-a3b-262k 生成自然语言回复。"],
            ["Embedding 服务", "http://127.0.0.1:8001/v1/embeddings", "调用 qwen3-embedding-8b 生成 4096 维语义向量。"],
            ["Web 入口", "0.0.0.0:9922 / 外部隧道 19922", "浏览器访问聊天页面，并通过 SSE 接收流式回复。"],
        ],
        [1.35, 2.25, 2.9],
    )

    doc.add_heading("3. 记忆工作流", level=1)
    add_numbered(
        doc,
        [
            "用户打开页面：后端创建新 session，前端显示空聊天；刷新或重新打开等价于 reset，但数据库不删除旧历史。",
            "用户发送消息：后端先写入 messages，再构造给主模型的 system prompt。",
            "查询向量化：后端调用 qwen3-embedding-8b，把当前用户问题转成 4096 维向量。",
            "相似检索：后端在 memory_vectors 中逐条计算余弦相似度，取 Top 10，再按 0.2 相似度阈值过滤。",
            "注入上下文：匹配到的历史片段会以“历史记忆 1/2/...”形式追加到 system prompt，主模型据此回答。",
            "回答完成：assistant 消息写入数据库；随后触发增量刷新，把新形成且尚无向量的历史片段补入 memory_vectors。",
        ],
    )

    doc.add_heading("4. 记忆切片与向量库原理", level=1)
    add_paragraph(
        doc,
        "系统不是把每条消息单独做 embedding，而是把同一 session 内的连续 user/assistant 消息组成滑动窗口片段。这样一个向量通常覆盖一小段上下文，而不是孤立的一句话，能保留问题、回答和前后语境。",
    )
    add_bullets(
        doc,
        [
            "窗口大小：默认 10 条 completed 的 user/assistant 消息组成一个 memory segment。",
            "滑动步长：默认 5 条消息，因此相邻片段有重叠，减少重要上下文被切断的概率。",
            "短会话处理：不足 10 条消息也会形成较短片段；当前数据库中的片段长度最小为 1，最大为 10，平均约 7.01 条消息。",
            "去重方式：片段文本会计算 content_hash，同样内容不会重复插入 memory_segments。",
            "向量保存：embedding 结果归一化后以 float32 BLOB 存入 SQLite，维度为 4096。",
        ],
    )

    doc.add_heading("5. 关键参数", level=1)
    params = metrics["params"]
    counts = metrics["counts"]
    embedding = metrics["embedding"]
    generation_model = metrics["generation_model"]
    segment_stats = metrics["segment_message_count"]
    add_kv_table(
        doc,
        [
            ["主模型", f"{generation_model['model']}，base_url={generation_model['base_url']}"],
            ["Embedding 模型", f"{embedding['model']}，base_url={embedding['base_url']}"],
            ["Embedding 维度", f"{embedding['vector_dim']} 维 float32 向量"],
            ["每段记忆包含多少消息", f"window_size={params['window_size']}；当前实际平均 {segment_stats['avg']:.2f} 条消息/片段"],
            ["滑动步长", f"stride={params['stride']}，相邻片段默认重叠 5 条消息"],
            ["每次检索 TopK", f"retrieval_top_k={params['retrieval_top_k']}"],
            ["相似度阈值", f"min_score={params['min_score']}；低于该阈值的片段不会注入 prompt"],
            ["每轮增量补向量数量", f"max_incremental_segments_per_turn={params['max_incremental_segments_per_turn']}"],
            ["当前历史规模", f"sessions={counts['sessions']}，messages={counts['messages']}，segments={counts['memory_segments']}，vectors={counts['memory_vectors']}"],
        ],
    )

    doc.add_heading("6. 实测性能", level=1)
    timing = metrics["timing_ms"]
    retrieval = metrics["retrieval_result"]
    add_paragraph(
        doc,
        "以下结果来自服务器当前状态下的 5 个中文样例问题。数值用于判断量级，不等价于长期压测；随着历史片段增长，暴力向量检索耗时会近似线性上升。",
    )
    add_matrix_table(
        doc,
        ["指标", "平均值", "P50", "说明"],
        [
            ["查询 embedding", f"{timing['embedding_avg']:.2f} ms", f"{timing['embedding_p50']:.2f} ms", "调用 qwen3-embedding-8b，把用户问题转向量。"],
            ["纯向量检索", f"{timing['retrieval_only_avg']:.2f} ms", f"{timing['retrieval_only_p50']:.2f} ms", "不含 embedding，仅从 SQLite 读取向量并计算余弦相似度。"],
            ["完整记忆构建", f"{timing['full_prompt_memory_avg']:.2f} ms", f"{timing['full_prompt_memory_p50']:.2f} ms", "包含查询 embedding、TopK 检索、上下文格式化。"],
            ["命中片段数", f"{retrieval['avg_matched_segments_after_threshold']:.1f}", "-", "阈值过滤后平均注入片段数。当前样例均达到 Top 10。"],
            ["注入上下文长度", f"{retrieval['avg_context_chars']:.0f} 字符", "-", "实际追加到 system prompt 的历史记忆文本长度。"],
        ],
        [1.45, 1.0, 0.9, 3.15],
    )

    doc.add_heading("7. 数据表与生命周期", level=1)
    add_matrix_table(
        doc,
        ["表", "内容", "生命周期"],
        [
            ["sessions", "会话 ID、访问 IP、User-Agent、开始/结束时间、结束原因。", "页面加载创建；reset/close 仅标记结束，不删除。"],
            ["messages", "用户和 assistant 的 completed/failed 消息。", "永久软保留，供历史查看和记忆构建使用。"],
            ["events", "session_start、message_user、vector_memory_error、vector_memory_refresh 等事件。", "用于排障、统计和追踪刷新状态。"],
            ["message_index / memory_fts", "旧版关键词检索索引。", "作为 fallback 使用。"],
            ["memory_segments", "连续聊天片段、起止 message_id、片段长度、content_hash。", "全量 rebuild 或每轮增量刷新产生。"],
            ["memory_vectors", "segment_id、向量维度、向量 BLOB、embedding 模型名。", "与 memory_segments 一一对应。"],
        ],
        [1.35, 2.75, 2.4],
    )

    doc.add_heading("8. 失败回退与边界", level=1)
    add_bullets(
        doc,
        [
            "Embedding 服务不可用时：本轮不会使用向量记忆，后端记录 vector_memory_error，并尝试关键词记忆。",
            "向量库为空时：不会注入历史记忆，主模型只使用基础 system prompt 和当前 session 历史。",
            "检索不是“读取所有历史”：每轮只注入最相似的约 10 段连续聊天，降低 prompt 膨胀。",
            "目前没有用户账号隔离：历史记忆会跨访客、跨 session 检索；这符合“让模型接触其他人历史”的目标，但不适合隐私敏感场景。",
            "SQLite 暴力检索适合当前百级/千级片段；当片段达到十万级，应迁移到 FAISS、Milvus、Qdrant 或 sqlite-vec。"
        ],
    )

    doc.add_heading("9. 运维命令", level=1)
    add_matrix_table(
        doc,
        ["任务", "命令"],
        [
            ["启动 Web 服务", "cd /base/home/lizhzh/Project3/qwen_web && ./start_qwen_web.sh"],
            ["停止 Web 服务", "cd /base/home/lizhzh/Project3/qwen_web && ./stop_qwen_web.sh"],
            ["启动 Embedding 服务", "cd /base/home/lizhzh/Project3/qwen_web && ./start_qwen_embedding.sh"],
            ["停止 Embedding 服务", "cd /base/home/lizhzh/Project3/qwen_web && ./stop_qwen_embedding.sh"],
            ["全量重建向量", "cd /base/home/lizhzh/Project3/qwen_web && /opt/conda/bin/python3 rebuild_memory_vectors.py --batch-size 4"],
            ["查看聊天记录", "cd /base/home/lizhzh/Project3/qwen_web && /opt/conda/bin/python3 data/view_chat_history.py"],
            ["健康检查", "curl http://127.0.0.1:9922/api/health"],
        ],
        [1.7, 4.8],
    )

    doc.add_heading("10. 后续建议", level=1)
    add_bullets(
        doc,
        [
            "增加记忆后台管理页：按时间、IP、session、相似度查看被检索到的片段。",
            "增加隐私隔离开关：支持全局记忆、同 IP 记忆、仅当前用户记忆三种模式。",
            "增加摘要层：把高频历史压缩为长期 persona / preference 记忆，减少重复加载原文。",
            "当 memory_segments 超过 1 万条时，引入近似向量索引，避免暴力检索线性增长。",
            "定期导出数据库备份，并对包含 IP 和聊天内容的 SQLite 文件设置访问权限。"
        ],
    )

    doc.add_section(WD_SECTION.NEW_PAGE)
    doc.add_heading("附录：本次采集口径", level=1)
    add_paragraph(doc, "采集时间：2026-05-28；采集对象：当前服务器 /base/home/lizhzh/Project3/qwen_web 部署。")
    add_paragraph(doc, "样本数量：5 个中文问题；耗时单位：毫秒；统计方法：Python time.perf_counter，平均值与 P50。")
    add_paragraph(doc, "注意：性能值受 GPU 负载、vLLM 批处理、SQLite 页缓存和当时并发影响，应视为当前量级参考。")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build_doc()
